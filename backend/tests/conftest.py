from __future__ import annotations

import os

# Быстрые тесты: без file watcher и фоновой индексации при старте API.
os.environ.setdefault("WATCH_DOCUMENTS", "false")
os.environ.setdefault("INDEX_ON_STARTUP", "false")

from pathlib import Path
import pytest
from docx import Document

FIXTURES_DIR = Path(__file__).parent / "fixtures"


@pytest.fixture
def sample_txt(tmp_path: Path) -> Path:
    path = tmp_path / "lecture.txt"
    path.write_text(
        "Введение в ML.\n\n"
        "Машинное обучение — подмножество AI.\n\n"
        "Нейронные сети состоят из слоёв.",
        encoding="utf-8",
    )
    return path


@pytest.fixture
def empty_txt(tmp_path: Path) -> Path:
    path = tmp_path / "empty.txt"
    path.write_text("   \n  ", encoding="utf-8")
    return path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    path = tmp_path / "notes.docx"
    doc = Document()
    doc.add_paragraph("Первый абзац лекции.")
    doc.add_paragraph("Второй абзац про нейросети.")
    doc.save(path)
    return path


@pytest.fixture
def sample_pdf() -> Path:
    return FIXTURES_DIR / "sample.pdf"
